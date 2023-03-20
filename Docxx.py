from docx import Document
from docx.shared import Inches
from docx import Document
import pandas as pd


df = pd.read_excel("dummy.xlsx")

df.drop(["Unnamed: 0"],axis=1,inplace=True)

df.drop(["RESULT_FORMATTED_ENTRY"],axis=1,inplace=True)
cc = df.columns

print("values are",df.values)

print("first value is",df.values[0])
print("secound value is",df.values[1])
cc = len(cc)

columm = list(df.columns)
temp_doc1 = Document("demo1.docx")


for table in temp_doc1.tables:
    print("table is",table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print("ccsdds",paragraph.text)

    for t in temp_doc1.tables:
        tab_cols=[]
        for col in t.columns:
            for c in col.cells:
                tab_cols.append(c.text)
                if tab_cols==columm:
                    for i in range(df.shape[0]):
                        row=t.add_row().cells
                        for j in range(df.shape[1]):
                            row[j].text = str(df.values[i,j])
                break


df = pd.read_excel("dummy.xlsx")

df.drop(["Unnamed: 0"],axis=1,inplace=True)

df.drop(["RESULT_FORMATTED_ENTRY"],axis=1,inplace=True)
cc = df.columns

cc = len(cc)

columm = list(df.columns)

print("columns are",columm[0])

#columm = ['Unn', 'CON', 'TIM', 'RES', 'App', 'App', 'Ass', 'Deg', 'Fin', 'Los', 'Mic', 'Uni', 'Uni']

document = Document()

document.add_heading('Document Title', 0)

#p = document.add_paragraph('A plain paragraph having some ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

#document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='Intense Quote')

#document.add_paragraph(
#    'first item in unordered list', style='List Bullet'
#)
#document.add_paragraph(
#    'first item in ordered list', style='List Number'
#)

#document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)


table1 = document.add_table(rows= 1,cols = 13)

hdr_cells = table1.rows[0].cells

ln = len(columm)

for i in range(0,11):
    hdr_cells[i].text = columm[i]

for index, row in df.iterrows():
    row_cells = table1.add_row().cells
    if index<=ln-1:
        print("index is",index)
        print("columm is:",row_cells[index])
        print("row is",row[columm[index]])
        row_cells[index].text = str(df.values[index,index])
        #row_cells[index].text = str(row[columm[index]])

document.add_page_break()

document.save('demo2.docx')


table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

