from docx import Document

temp_doc =Document("Sample.docx")

print(temp_doc.tables)
print(temp_doc.paragraphs)

for table in temp_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)


# for i in temp_doc.paragraphs:
#     print("paragraph is",i.text)
